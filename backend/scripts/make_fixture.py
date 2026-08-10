"""Generate the smoke-test resume fixture.

    docker compose run --rm --no-deps -T api python scripts/make_fixture.py

A real .docx rather than text-with-a-.pdf-extension, because the parser
correctly rejects the latter and a smoke test that uploads a fake file only
proves the rejection path works.
"""

from __future__ import annotations

from pathlib import Path

import docx

OUT = Path(__file__).parent / "fixtures" / "sample-resume.docx"

LINES = [
    ("Alex Candidate", None),
    ("alex@example.com | +44 7700 900123", None),
    ("Experience", "Heading 1"),
    (
        "Senior Backend Engineer, Orderly (2021-2026). Owned indexing strategy and "
        "query planning for the orders service, cutting p95 read latency by half.",
        "List Bullet",
    ),
    (
        "Designed the rest api design and error contract design for the public API, "
        "including api versioning and a deprecation policy.",
        "List Bullet",
    ),
    (
        "Introduced idempotency keys and rate limiting across the payment endpoints "
        "after a duplicate-charge incident.",
        "List Bullet",
    ),
    (
        "Led connection pooling and schema migration safety work during the Postgres "
        "upgrade, with zero downtime.",
        "List Bullet",
    ),
    ("Projects", "Heading 1"),
    (
        "Ledger: an event-sourced balances service using the outbox pattern and "
        "at-least-once and idempotent consumers.",
        "List Bullet",
    ),
    ("Skills", "Heading 1"),
    (
        "Python, PostgreSQL, transactions and acid, isolation levels and anomalies, "
        "caching strategies, background jobs and queues, async concurrency model",
        None,
    ),
]


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    document = docx.Document()
    for text, style in LINES:
        document.add_paragraph(text, style=style) if style else document.add_paragraph(text)
    document.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
